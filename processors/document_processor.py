import io
import re
import hashlib
import json
import datetime
from typing import Dict, Any, List, Optional, Tuple
from collections import defaultdict

# file handling libraries
import fitz  # PyMuPDF
from PIL import Image, ImageOps
import pytesseract
from docx import Document
import piexif
import filetype
import tldextract

# NER: spaCy
import spacy

# helper libs
import base64
import math

# Try load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except Exception:
    # fallback: try to load English transformer if available
    try:
        nlp = spacy.load("en_core_web_trf")
    except Exception:
        nlp = None


# --- Utilities -----------------------------------------------------------------
def now_iso() -> str:
    return datetime.datetime.utcnow().replace(microsecond=0).isoformat() + "Z"


def sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


def safe_text(s) -> Optional[str]:
    if s is None:
        return None
    s = str(s).strip()
    return s if s else None


# Regex detectors
RE_EMAIL = re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+")
RE_URL = re.compile(r"(https?://[^\s'\"<>]+)")
RE_IP = re.compile(r"\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d?\d)\.){3}(?:25[0-5]|2[0-4]\d|[01]?\d?\d)\b")
RE_PHONE = re.compile(r"(?:\+?\d{1,3}[-.\s]?)?(?:\(?\d{2,4}\)?[-.\s]?)?\d{3,4}[-.\s]?\d{3,4}")


# --- Extractors ----------------------------------------------------------------
class TextBlock:
    def __init__(self, text: str, bbox: Tuple[int, int, int, int], conf: float = 1.0):
        self.text = text
        self.bbox = bbox  # x,y,w,h
        self.confidence = conf


class DocumentProcessor:
    """
    Usage:
      DocumentProcessor.process(file_bytes, file_extension)
    file_extension: 'pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg', 'eml', etc.
    """

    SUPPORTED_IMAGE_FORMATS = {"png", "jpg", "jpeg", "tiff", "bmp", "gif"}

    @classmethod
    def process(cls, file_data: bytes, file_extension: str) -> Dict[str, Any]:
        ext = file_extension.lower().lstrip(".")
        result = {
            "metadata": {
                "document_title": None,
                "author": None,
                "creation_date": None,
                "mod_date": None,
                "last_saved_by": None,
                "application": None,
                "page_count": None,
                "language": None,
                "fonts": [],
            },
            "structure": {
                "pages": [],
                "document_text": "",
                "char_count": 0
            },
            "entities": {
                "named_entities": [],
                "pii_summary": {"emails": [], "phones": [], "ids": []}
            },
            "indicators": {
                "urls": [],
                "domains": [],
                "ips": [],
                "file_hashes": [],
                "emails": []
            },
            "image_analysis": {
                "image_count": 0,
                "exif": [],
                "faces_detected": [],
                "steganography_suspect": []
            },
            "nlp": {
                "language": "und",
                "summary": None,
                "keywords": [],
                "sentiment": {"polarity": 0.0, "subjectivity": 0.0},
                "embedding_id": None
            },
            "feature_vector": {
                "numeric_features": {
                    "word_count": 0,
                    "unique_urls": 0,
                    "suspicious_url_count": 0,
                    "pii_count": 0,
                    "yara_count": 0,
                    "avg_entropy": 0.0
                },
                "dense_vector_reference": None
            },
            "confidence": {
                "overall_confidence": 0.0,
                "component_confidence": {
                    "ocr": 0.0,
                    "metadata": 0.0,
                    "entity_extraction": 0.0,
                    "image_analysis": 0.0,
                    "binary_analysis": 0.0
                }
            }
        }

        # add file hashes
        result["indicators"]["file_hashes"].append({"alg": "sha256", "value": sha256_bytes(file_data), "source": "input_blob"})
        # detect type heuristics
        if ext in ("pdf",):
            cls._process_pdf(file_data, result)
        elif ext in ("docx", "doc"):
            cls._process_docx(file_data, result)
        elif ext in ("txt", "text"):
            cls._process_txt(file_data, result)
        elif ext in cls.SUPPORTED_IMAGE_FORMATS:
            cls._process_image(file_data, result, ext)
        else:
            # try auto-detect via filetype
            kind = filetype.guess(file_data)
            if kind:
                if kind.mime.startswith("image/"):
                    cls._process_image(file_data, result, kind.extension)
                elif kind.extension == "pdf":
                    cls._process_pdf(file_data, result)
                else:
                    # fallback plain text try
                    cls._process_txt(file_data, result)

        # Run entity extraction / indicators on document_text
        doc_text = result["structure"]["document_text"]
        cls._extract_indicators_and_entities(doc_text, result)

        # Basic NLP summary/keywords using spaCy (if available)
        cls._basic_nlp(result)

        # finalize numeric features
        cls._finalize_features(result)

        # compute overall confidence as weighted mean of component confidences
        comps = result["confidence"]["component_confidence"]
        weights = {"ocr": 0.25, "metadata": 0.15, "entity_extraction": 0.3, "image_analysis": 0.2, "binary_analysis": 0.1}
        total_weight = sum(weights.values())
        overall = sum(comps.get(k, 0.0) * w for k, w in weights.items()) / total_weight
        result["confidence"]["overall_confidence"] = round(min(max(overall, 0.0), 1.0), 3)

        return result

    # ---------------- PDF ----------------
    @classmethod
    def _process_pdf(cls, file_data: bytes, result: Dict[str, Any]):
        try:
            
            doc = fitz.Document(stream=file_data, filetype="pdf")
        except Exception as e:
            # can't open PDF - leave result mostly empty but set binary_analysis confidence low
            result["confidence"]["component_confidence"]["binary_analysis"] = 0.05
            return

        # metadata
        meta = doc.metadata or {}
        result["metadata"]["document_title"] = safe_text(meta.get("title"))
        result["metadata"]["author"] = safe_text(meta.get("author"))
        # fitz stores dates like "D:YYYYMMDDHHmmSS"
        def parse_fit_time(v):
            if not v:
                return None
            try:
                # quick parse for common pdf format D:YYYYMMDDHHmmSS
                m = re.match(r"D:(\d{4})(\d{2})?(\d{2})?(\d{2})?(\d{2})?(\d{2})?.*", v)
                if m:
                    parts = m.groups()
                    year = int(parts[0])
                    month = int(parts[1] or 1)
                    day = int(parts[2] or 1)
                    return datetime.datetime(year, month, day).isoformat() + "Z"
            except Exception:
                pass
            return None

        result["metadata"]["creation_date"] = parse_fit_time(meta.get("creationDate"))
        result["metadata"]["mod_date"] = parse_fit_time(meta.get("modDate"))
        result["metadata"]["application"] = safe_text(meta.get("producer"))
        result["metadata"]["page_count"] = doc.page_count

        pages_text_parts = []
        ocr_conf_agg = []
        image_count = 0
        fonts_set = set()

        for pno in range(doc.page_count):
            page = doc.load_page(pno)
            page_text_dict = page.get_text("dict")
            if isinstance(page_text_dict, dict):
                blocks = page_text_dict.get("blocks", [])
            else:
                blocks = []
            page_entry = {"page_number": pno + 1, "text_blocks": [], "tables": [], "images": [], "annotations": []}

            # text / bbox extraction using PyMuPDF text blocks
            for b in blocks:
                if b.get("type") == 0:  # text block
                    for line in b.get("lines", []):
                        txt = "".join([span["text"] for span in line.get("spans", [])])
                        # bounding box approx
                        bbox = line.get("bbox", [0, 0, 0, 0])
                        x0, y0, x1, y1 = bbox
                        w, h = int(x1 - x0), int(y1 - y0)
                        tb = {"id": f"p{pno+1}_t{len(page_entry['text_blocks'])+1}",
                              "bbox": [int(x0), int(y0), w, h],
                              "text": txt,
                              "confidence": 1.0}
                        page_entry["text_blocks"].append(tb)
                        pages_text_parts.append(txt)

                        # collect spans for fonts
                        for span in line.get("spans", []):
                            if span.get("font"):
                                fonts_set.add(span.get("font"))

                elif b.get("type") == 1:  # image block
                    # extract image bytes
                    for img in b.get("image", []):
                        pass
                    # simpler: render section to an image and OCR
                    try:
                        pix = page.get_pixmap()
                        img_bytes = pix.tobytes()
                        img_sha = sha256_bytes(img_bytes)
                        image_count += 1
                        # OCR the whole page image as fallback for scanned PDFs
                        pil = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                        ocr_text = cls._run_ocr_on_pil(pil)
                        img_id = f"p{pno+1}_img{image_count}"
                        page_entry["images"].append({
                            "id": img_id,
                            "bbox": [0, 0, pix.width, pix.height],
                            "sha256": img_sha,
                            "mime_type": "image/png",
                            "ocr_text": ocr_text
                        })
                        ocr_conf_agg.append(0.9)  # heuristic
                    except Exception:
                        pass

            # Append page entry
            result["structure"]["pages"].append(page_entry)

        # finalize
        result["structure"]["document_text"] = "\n".join(pages_text_parts).strip()
        result["structure"]["char_count"] = len(result["structure"]["document_text"])
        result["image_analysis"]["image_count"] = image_count
        result["metadata"]["fonts"] = sorted(list(fonts_set))[:200]

        # set confidence heuristics
        result["confidence"]["component_confidence"]["metadata"] = 0.9 if result["metadata"]["page_count"] else 0.4
        # ocr: if text extracted length is small but pages exist and images exist, lower OCR confidence
        if result["structure"]["char_count"] > 100:
            result["confidence"]["component_confidence"]["ocr"] = 0.95
        else:
            result["confidence"]["component_confidence"]["ocr"] = 0.6

        result["confidence"]["component_confidence"]["image_analysis"] = 0.7 if image_count > 0 else 0.2
        result["confidence"]["component_confidence"]["binary_analysis"] = 0.8

    # ---------------- DOCX ----------------
    @classmethod
    def _process_docx(cls, file_data: bytes, result: Dict[str, Any]):
        try:
            doc = Document(io.BytesIO(file_data))
        except Exception:
            result["confidence"]["component_confidence"]["binary_analysis"] = 0.1
            return

        # metadata via core_properties
        core = doc.core_properties
        result["metadata"]["document_title"] = safe_text(core.title)
        result["metadata"]["author"] = safe_text(core.author)
        if core.created:
            try:
                result["metadata"]["creation_date"] = core.created.isoformat() + "Z"
            except Exception:
                result["metadata"]["creation_date"] = safe_text(core.created)
        if core.modified:
            try:
                result["metadata"]["mod_date"] = core.modified.isoformat() + "Z"
            except Exception:
                result["metadata"]["mod_date"] = safe_text(core.modified)

        pages_text_parts = []
        page_index = 1
        # python-docx doesn't expose page layout; treat whole doc as single logical page series
        page_entry = {"page_number": 1, "text_blocks": [], "tables": [], "images": [], "annotations": []}
        tbl_count = 0
        img_count = 0
        for block in doc.element.body:
            tag = getattr(block, "tag", "")
            text = ""
            try:
                text = block.text if hasattr(block, "text") else ""
            except Exception:
                text = ""
            if text:
                tb = {"id": f"p{page_index}_t{len(page_entry['text_blocks'])+1}", "bbox": [0, 0, 0, 0], "text": text, "confidence": 1.0}
                page_entry["text_blocks"].append(tb)
                pages_text_parts.append(text)

        # tables
        for tbl in doc.tables:
            tbl_count += 1
            rows = len(tbl.rows)
            cols = len(tbl.columns)
            # produce csv
            cells = []
            for r in range(rows):
                row_vals = []
                for c in range(cols):
                    try:
                        row_vals.append(tbl.cell(r, c).text.replace("\n", " ").strip())
                    except Exception:
                        row_vals.append("")
                cells.append(",".join(['"{}"'.format(v.replace('"', '""')) for v in row_vals]))
            csv_text = "\n".join(cells)
            page_entry["tables"].append({
                "id": f"tbl{tbl_count}",
                "bbox": [0, 0, 0, 0],
                "rows": rows,
                "cols": cols,
                "extracted_csv": csv_text
            })

        # images: python-docx image extraction is non-trivial; we can scan the package media parts if needed.
        # For now, count 0 or implement via doc.part.related_parts (left as an exercise for clarity)
        result["structure"]["pages"].append(page_entry)
        result["structure"]["document_text"] = "\n".join(pages_text_parts).strip()
        result["structure"]["char_count"] = len(result["structure"]["document_text"])
        result["image_analysis"]["image_count"] = img_count

        # confidences
        result["confidence"]["component_confidence"]["metadata"] = 0.9
        result["confidence"]["component_confidence"]["ocr"] = 0.0  # OCR not used for docx
        result["confidence"]["component_confidence"]["image_analysis"] = 0.5 if img_count else 0.2
        result["confidence"]["component_confidence"]["binary_analysis"] = 0.9

    # ---------------- TXT ----------------
    @classmethod
    def _process_txt(cls, file_data: bytes, result: Dict[str, Any]):
        try:
            txt = file_data.decode("utf-8", errors="replace")
        except Exception:
            txt = str(file_data)
        lines = txt.splitlines()
        page_entry = {"page_number": 1, "text_blocks": [], "tables": [], "images": [], "annotations": []}
        for idx, line in enumerate(lines):
            tb = {"id": f"p1_t{idx+1}", "bbox": [0, 0, 0, 0], "text": line, "confidence": 1.0}
            page_entry["text_blocks"].append(tb)
        result["structure"]["pages"].append(page_entry)
        result["structure"]["document_text"] = txt
        result["structure"]["char_count"] = len(txt)
        result["confidence"]["component_confidence"]["metadata"] = 0.1
        result["confidence"]["component_confidence"]["ocr"] = 0.0
        result["confidence"]["component_confidence"]["entity_extraction"] = 0.8
        result["confidence"]["component_confidence"]["binary_analysis"] = 0.6

    # ---------------- Image ----------------
    @classmethod
    def _process_image(cls, file_data: bytes, result: Dict[str, Any], ext: str):
        try:
            from processors.image_processor import ImageProcessor
            image_report = ImageProcessor.process(file_data, f'.{ext}')
            result['image_analysis']['image_processor_output'] = image_report
        except Exception:
            result["confidence"]["component_confidence"]["binary_analysis"] = 0.1
            return
        # Optionally, set some basic confidences
        result["confidence"]["component_confidence"]["ocr"] = 0.85
        result["confidence"]["component_confidence"]["image_analysis"] = 0.8
        result["confidence"]["component_confidence"]["binary_analysis"] = 0.7

    # ---------------- OCR helper ----------------
    @classmethod
    def _run_ocr_on_pil(cls, pil_img: Image.Image) -> str:
        try:
            # convert to grayscale and enhance
            gray = ImageOps.grayscale(pil_img)
            # optional: thresholding could be done for better OCR
            txt = pytesseract.image_to_string(gray)
            return txt.strip()
        except Exception:
            return ""

    # ---------------- indicators & entities ----------------
    @classmethod
    def _extract_indicators_and_entities(cls, doc_text: str, result: Dict[str, Any]):
        if not doc_text:
            result["confidence"]["component_confidence"]["entity_extraction"] = 0.05
            return

        # quick regex-based PII
        emails = sorted(set(RE_EMAIL.findall(doc_text)))
        phones = sorted(set(RE_PHONE.findall(doc_text)))
        ips = sorted(set(RE_IP.findall(doc_text)))
        urls = sorted(set(RE_URL.findall(doc_text)))

        # populate indicators arrays
        for u in urls:
            parsed = tldextract.extract(u)
            hostname = ".".join([p for p in (parsed.subdomain, parsed.domain, parsed.suffix) if p])
            result["indicators"]["urls"].append({
                "url": u,
                "hostname": hostname,
                "port": None,
                "path": "",
                "query": None,
                "is_ip": bool(RE_IP.search(u)),
                "suspicious_score": 0.0  # placeholder; call reputation service to fill
            })
        # domains (unique)
        domain_set = set([d["hostname"] for d in result["indicators"]["urls"] if d["hostname"]])
        for d in domain_set:
            result["indicators"]["domains"].append({"domain": d, "whois": {"registrar": None, "created": None}, "reputation_score": 0.0})

        for ip in ips:
            result["indicators"]["ips"].append({"ip": ip, "asn": None, "geo": None, "reputation_score": 0.0})

        for e in emails:
            result["entities"]["pii_summary"]["emails"].append(e)
            result["indicators"]["emails"].append({"addr": e, "display_name": None, "header_excerpt": None})

        for p in phones:
            result["entities"]["pii_summary"]["phones"] = result["entities"]["pii_summary"].get("phones", []) + [p]

        # file hashes already added earlier

        # NER via spaCy if available
        ner_conf = 0.0
        if nlp:
            try:
                doc = nlp(doc_text)
                for ent in doc.ents:
                    typ = ent.label_
                    mapped = cls._map_spacy_label(typ)
                    ne = {
                        "type": mapped,
                        "text": ent.text,
                        "start_char": ent.start_char,
                        "end_char": ent.end_char,
                        "confidence": getattr(ent, "kb_id_", 0.9) or 0.9
                    }
                    result["entities"]["named_entities"].append(ne)
                ner_conf = 0.9
            except Exception:
                ner_conf = 0.4
        else:
            # Basic heuristics: add emails/ips/urls as named entities
            for e in emails:
                result["entities"]["named_entities"].append({"type": "EMAIL", "text": e, "start_char": doc_text.find(e), "end_char": doc_text.find(e) + len(e), "confidence": 0.9})
            for ip in ips:
                result["entities"]["named_entities"].append({"type": "IP", "text": ip, "start_char": doc_text.find(ip), "end_char": doc_text.find(ip) + len(ip), "confidence": 0.9})
            ner_conf = 0.6

        result["confidence"]["component_confidence"]["entity_extraction"] = ner_conf

    @staticmethod
    def _map_spacy_label(label: str) -> str:
        # Map various spaCy labels to requested schema
        m = {
            "PERSON": "PERSON", "ORG": "ORG", "GPE": "GPE", "LOC": "LOCATION", "DATE": "DATE", "TIME": "DATE",
            "MONEY": "OTHER", "PERCENT": "OTHER", "EMAIL": "EMAIL", "PHONE": "PHONE", "IP": "IP", "CARDINAL": "OTHER",
        }
        return m.get(label, "OTHER")

    # ---------------- NLP summary, keywords (basic) ----------------
    @classmethod
    def _basic_nlp(cls, result: Dict[str, Any]):
        text = result["structure"]["document_text"]
        if not text:
            result["nlp"]["language"] = "und"
            result["nlp"]["summary"] = None
            result["nlp"]["keywords"] = []
            return
        # language detection: naive English detection via spaCy presence
        result["nlp"]["language"] = "en" if (nlp is not None) else "und"
        # Basic summary: first 3 non-empty lines or first 300 chars
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        if lines:
            result["nlp"]["summary"] = " ".join(lines[:3])[:800]
        else:
            result["nlp"]["summary"] = text[:800]
        # keywords: top nouns / proper nouns via spaCy if available
        keywords = []
        if nlp:
            try:
                doc = nlp(text[:5000])  # limit window
                freq = defaultdict(int)
                for tok in doc:
                    if tok.pos_ in ("NOUN", "PROPN") and len(tok.text) > 2:
                        freq[tok.lemma_.lower()] += 1
                keywords = sorted(freq.items(), key=lambda x: -x[1])[:20]
                keywords = [k for k, _ in keywords]
            except Exception:
                keywords = []
        else:
            # fallback: top words by regex
            words = re.findall(r"\b[a-zA-Z]{3,}\b", text.lower())
            freq = defaultdict(int)
            for w in words:
                freq[w] += 1
            keywords = sorted(freq.items(), key=lambda x: -x[1])[:20]
            keywords = [k for k, _ in keywords]
        result["nlp"]["keywords"] = keywords

    # ---------------- finalize feature vector ----------------
    @classmethod
    def _finalize_features(cls, result: Dict[str, Any]):
        txt = result["structure"]["document_text"]
        words = re.findall(r"\w+", txt)
        word_count = len(words)
        unique_urls = len({u["hostname"] for u in result["indicators"]["urls"]})
        suspicious_url_count = sum(1 for u in result["indicators"]["urls"] if u.get("suspicious_score", 0) > 0.5)
        pii_count = len(result["entities"]["pii_summary"].get("emails", [])) + len(result["entities"]["pii_summary"].get("phones", []))
        # avg entropy per line heuristic
        entropies = []
        for s in txt.splitlines():
            if not s:
                continue
            freq = defaultdict(int)
            for ch in s:
                freq[ch] += 1
            l = len(s)
            ent = -sum((v/l)*math.log2(v/l) for v in freq.values() if v > 0)
            entropies.append(ent)
        avg_entropy = sum(entropies) / len(entropies) if entropies else 0.0

        result["feature_vector"]["numeric_features"].update({
            "word_count": word_count,
            "unique_urls": unique_urls,
            "suspicious_url_count": suspicious_url_count,
            "pii_count": pii_count,
            "yara_count": 0,
            "avg_entropy": round(avg_entropy, 3)
        })

        # dense_vector_reference left as None; call out to vector-db ingestion in production flow
        result["feature_vector"]["dense_vector_reference"] = None

        # counts used elsewhere
        result["nlp"]["embedding_id"] = None

        # set final confidence defaults (if not already set)
        for k in ["ocr", "metadata", "entity_extraction", "image_analysis", "binary_analysis"]:
            if result["confidence"]["component_confidence"].get(k) is None:
                result["confidence"]["component_confidence"][k] = 0.5

